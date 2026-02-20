#!/usr/bin/env python3
"""
FAQ Generator - Clean DOCX Output

Generates a comprehensive Ukrainian FAQ document from Signal messages.
- Uses Gemini to extract and cluster questions/answers
- Tracks references internally for quality assurance
- Outputs clean DOCX without visible references

Usage:
    python scripts/generate_faq_docx.py

Output:
    docs/FAQ.docx - Clean Ukrainian FAQ for users
"""

import os
import sys
import json
import time
from pathlib import Path
from datetime import datetime
from collections import Counter

# Fix encoding for Windows console
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')

def _maybe_load_dotenv(dotenv_path):
    """Load key=value pairs from .env."""
    path = Path(dotenv_path)
    if not path.exists():
        return
    for raw in path.read_text(encoding="utf-8", errors="ignore").splitlines():
        line = raw.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, v = line.split("=", 1)
        k = k.strip()
        v = v.strip().strip("\r")
        if not k:
            continue
        if (v.startswith("'") and v.endswith("'")) or (v.startswith('"') and v.endswith('"')):
            v = v[1:-1]
        os.environ.setdefault(k, v)

# Load .env
_maybe_load_dotenv(Path(__file__).parent.parent / ".env")
_maybe_load_dotenv(".env")

from google import genai
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Configuration
GOOGLE_API_KEY = os.environ.get("GOOGLE_API_KEY")
if not GOOGLE_API_KEY:
    raise ValueError("GOOGLE_API_KEY not set")

client = genai.Client(api_key=GOOGLE_API_KEY)
MODEL_NAME = "gemini-2.0-flash"
BATCH_SIZE = 200
MAX_BATCHES = 15


def load_messages(filepath: str) -> list:
    """Load messages from JSON file."""
    print(f"Loading messages from {filepath}...")
    with open(filepath, 'r', encoding='utf-8') as f:
        data = json.load(f)
    messages = data.get('messages', [])
    print(f"Loaded {len(messages)} messages")
    return messages


def filter_relevant_messages(messages: list) -> list:
    """Filter out system messages, keep meaningful content."""
    relevant = []
    for msg in messages:
        if msg.get('type') in ['group-v2-change', 'keychange', 'verified-change']:
            continue
        body = msg.get('body', '').strip()
        if not body:
            continue
        if body.startswith('[ATTACHMENT') and '\n' not in body and body.endswith(']'):
            continue
        relevant.append({
            'body': body,
            'type': msg.get('type'),
            'timestamp': msg.get('timestamp'),
            'id': msg.get('id', '')
        })
    print(f"Filtered to {len(relevant)} relevant messages")
    return relevant


def chunk_messages(messages: list, batch_size: int = BATCH_SIZE) -> list:
    """Split messages into batches."""
    batches = []
    for i in range(0, len(messages), batch_size):
        batches.append(messages[i:i + batch_size])
    return batches


def extract_qa_from_batch(batch: list, batch_num: int, global_offset: int) -> str:
    """Extract Q&A pairs from a batch of messages."""
    
    messages_text = "\n---\n".join([
        f"[MSG_{global_offset + i + 1}] {msg['body']}" 
        for i, msg in enumerate(batch)
    ])
    
    prompt = f"""Проаналізуй ці повідомлення з чату технічної підтримки ArduPilot/StabX/дронів.

ЗАВДАННЯ: Виділи пари ПИТАННЯ-ВІДПОВІДЬ які реально присутні в чаті.

Повідомлення (batch {batch_num}):
{messages_text}

ПРАВИЛА:
1. Виділяй тільки РЕАЛЬНІ питання що були задані в чаті
2. Виділяй тільки РЕАЛЬНІ відповіді що були надані
3. Зберігай технічні терміни та параметри точно як в оригіналі
4. Вказуй номери повідомлень для кожного Q і A

Формат JSON:
{{
    "qa_pairs": [
        {{
            "question": "Текст питання українською",
            "answer": "Текст відповіді українською", 
            "q_refs": ["MSG_123"],
            "a_refs": ["MSG_124", "MSG_125"],
            "topic": "категорія (Налаштування/Проблеми/Камери/Живлення/тощо)"
        }}
    ],
    "standalone_questions": [
        {{"q": "Питання без відповіді", "refs": ["MSG_200"]}}
    ]
}}

Поверни тільки валідний JSON."""

    try:
        response = client.models.generate_content(model=MODEL_NAME, contents=prompt)
        return response.text
    except Exception as e:
        print(f"Error batch {batch_num}: {e}")
        return "{}"


def merge_qa_data(all_extracts: list) -> dict:
    """Merge Q&A data from all batches."""
    merged = {
        'qa_pairs': [],
        'standalone_questions': []
    }
    
    for extract_text in all_extracts:
        try:
            text = extract_text.strip()
            if text.startswith('```'):
                text = text.split('\n', 1)[1]
            if text.endswith('```'):
                text = text.rsplit('```', 1)[0]
            text = text.strip()
            
            data = json.loads(text)
            merged['qa_pairs'].extend(data.get('qa_pairs', []))
            merged['standalone_questions'].extend(data.get('standalone_questions', []))
        except (json.JSONDecodeError, Exception) as e:
            print(f"Warning: Parse error: {e}")
            continue
    
    return merged


def deduplicate_and_rank(merged: dict) -> list:
    """Deduplicate similar Q&A pairs and rank by frequency."""
    
    # Group similar questions
    question_groups = {}
    
    for qa in merged['qa_pairs']:
        q = qa.get('question', '').strip()
        if not q:
            continue
        
        # Normalize for grouping
        normalized = ' '.join(q.lower().split())[:100]
        
        if normalized not in question_groups:
            question_groups[normalized] = {
                'questions': [],
                'answers': [],
                'refs': [],
                'topics': []
            }
        
        question_groups[normalized]['questions'].append(q)
        if qa.get('answer'):
            question_groups[normalized]['answers'].append(qa['answer'])
        question_groups[normalized]['refs'].extend(qa.get('q_refs', []))
        question_groups[normalized]['refs'].extend(qa.get('a_refs', []))
        if qa.get('topic'):
            question_groups[normalized]['topics'].append(qa['topic'])
    
    # Rank by frequency and create final list
    ranked = []
    for norm_q, data in question_groups.items():
        freq = len(data['questions'])
        if freq >= 1:  # Include questions that appear at least once with answer
            # Pick best question text (longest/most detailed)
            best_q = max(data['questions'], key=len)
            # Pick best answer (longest/most detailed)
            best_a = max(data['answers'], key=len) if data['answers'] else None
            # Most common topic
            topic = Counter(data['topics']).most_common(1)[0][0] if data['topics'] else "Інше"
            
            if best_a:  # Only include if we have an answer
                ranked.append({
                    'question': best_q,
                    'answer': best_a,
                    'frequency': freq,
                    'refs': list(set(data['refs'])),
                    'topic': topic
                })
    
    # Sort by frequency (most frequent first)
    ranked.sort(key=lambda x: x['frequency'], reverse=True)
    return ranked


def generate_comprehensive_faq(ranked_qa: list, total_messages: int) -> list:
    """Use Gemini to generate comprehensive FAQ from ranked Q&A pairs."""
    
    # Prepare data with refs for quality control - only include Q&A with substantial answers
    qa_items = []
    for qa in ranked_qa[:80]:
        answer = qa.get('answer', '')
        # Skip if answer is too short or generic
        if len(answer) < 30:
            continue
        # Skip generic/useless answers
        skip_patterns = [
            'уточніть', 'залежить від', 'це рекомендація', 'від чого', 
            'може бути', 'спробуйте', 'перевірте', 'можливо'
        ]
        answer_lower = answer.lower()
        # Only skip if the ENTIRE answer is just one of these generic phrases
        if len(answer) < 60 and any(answer_lower.strip().startswith(p) for p in skip_patterns):
            continue
            
        refs = qa.get('refs', [])[:3]
        refs_str = f" [refs: {', '.join(refs)}]" if refs else ""
        qa_items.append(
            f"[{qa['frequency']}x | {qa['topic']}]\n"
            f"Q: {qa['question']}\n"
            f"A: {qa['answer']}{refs_str}"
        )
    
    qa_summary = "\n\n".join(qa_items[:50])
    
    prompt = f"""Створи ПРОФЕСІЙНИЙ FAQ документ на основі реальних питань-відповідей з чату підтримки.

ВХІДНІ ДАНІ (реальні Q&A з чату, [refs] = посилання на повідомлення):
{qa_summary}

КРИТИЧНІ ПРАВИЛА:

1. ПИТАННЯ - ПЕРЕФОРМУЛЮЙ ПРОФЕСІЙНО:
   - Видали ВСІ привітання ("Вітаю", "Привіт", "Доброго", "Панове", "друзі")
   - Видали звернення ("підкажіть", "допоможіть", "порадьте", "скажіть")
   - Зроби питання ЗАГАЛЬНИМИ (без "мій дрон", "у мене")
   - ПОГАНО: "Вітаю, підкажіть яку камеру краще обрати для мого дрона?"
   - ДОБРЕ: "Які камери рекомендовані для StabX?"

2. ВІДПОВІДІ - ТІЛЬКИ КОНКРЕТНІ:
   - КОЖНА відповідь МУСИТЬ містити КОНКРЕТНУ інформацію:
     * Назви параметрів (BRD_ALT_CONFIG, FS_EKF_THRESH, тощо)
     * Конкретні значення (= 1, = 230400, >= 1.5)
     * Назви моделей (Matek H743, OV5647-120, Caddx256)
     * Конкретні дії (зайти на 5050, натиснути Download)
   - ВИКЛЮЧИ якщо відповідь не містить конкретики
   - ВИКЛЮЧИ: "Необхідно встановити параметр X" (без значення)
   - ВИКЛЮЧИ: "Перевірте налаштування" (без деталей)
   - ВИКЛЮЧИ: "Залежить від ситуації"

3. НЕ ВИГАДУЙ:
   - Використовуй ТІЛЬКИ інформацію з даних вище
   - Якщо значення параметра не вказано в даних - НЕ ДОДАВАЙ це питання
   - Краще менше питань, але всі з корисними відповідями

ФОРМАТ - 15-25 FAQ записів:

{{
    "faq": [
        {{
            "category": "Категорія",
            "question": "Чисте професійне питання?",
            "answer": "Конкретна відповідь з параметрами/значеннями/моделями."
        }}
    ]
}}

Категорії: Налаштування ArduPilot, Інтеграція StabX, Камери та відео, Живлення, Проблеми та вирішення, Прошивка та оновлення, Ліцензування

Поверни тільки валідний JSON."""

    try:
        response = client.models.generate_content(model=MODEL_NAME, contents=prompt)
        text = response.text.strip()
        if text.startswith('```'):
            text = text.split('\n', 1)[1]
        if text.endswith('```'):
            text = text.rsplit('```', 1)[0]
        data = json.loads(text.strip())
        return data.get('faq', [])
    except Exception as e:
        print(f"Error generating FAQ: {e}")
        return []


def create_docx(faq_items: list, output_path: str, total_messages: int):
    """Create a professional DOCX document."""
    
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading('FAQ: Підтримка ArduPilot та StabX', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacer
    
    # Group by category
    categories = {}
    for item in faq_items:
        cat = item.get('category', 'Інше')
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(item)
    
    # Add each category
    for category, items in categories.items():
        # Category heading
        doc.add_heading(category, level=1)
        
        for item in items:
            # Question (bold)
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"Q: {item.get('question', '')}")
            q_run.bold = True
            q_run.font.size = Pt(11)
            
            # Answer
            a_para = doc.add_paragraph()
            a_run = a_para.add_run(f"A: {item.get('answer', '')}")
            a_run.font.size = Pt(11)
            
            # Small spacer
            doc.add_paragraph()
    
    # Save
    doc.save(output_path)
    print(f"Saved DOCX to {output_path}")


def main():
    print("=" * 60)
    print("FAQ Generator - DOCX Output")
    print("=" * 60)
    
    repo_root = Path(__file__).parent.parent
    data_file = repo_root / "test" / "data" / "signal_messages.json"
    
    if not data_file.exists():
        print(f"ERROR: Data file not found: {data_file}")
        sys.exit(1)
    
    # Load and filter
    messages = load_messages(str(data_file))
    relevant = filter_relevant_messages(messages)
    
    # Batch processing
    batches = chunk_messages(relevant, BATCH_SIZE)
    batches = batches[:MAX_BATCHES]
    print(f"Processing {len(batches)} batches...")
    
    # Extract Q&A from each batch
    all_extracts = []
    global_offset = 0
    for i, batch in enumerate(batches):
        print(f"\nBatch {i+1}/{len(batches)} ({len(batch)} messages)...")
        extract = extract_qa_from_batch(batch, i+1, global_offset)
        all_extracts.append(extract)
        global_offset += len(batch)
        time.sleep(1)
    
    # Merge and deduplicate
    print("\nMerging Q&A data...")
    merged = merge_qa_data(all_extracts)
    print(f"  - Q&A pairs: {len(merged['qa_pairs'])}")
    print(f"  - Standalone questions: {len(merged['standalone_questions'])}")
    
    # Rank by frequency
    print("\nDeduplicating and ranking...")
    ranked = deduplicate_and_rank(merged)
    print(f"  - Unique Q&A: {len(ranked)}")
    print(f"  - Frequent (2+): {len([r for r in ranked if r['frequency'] >= 2])}")
    
    # Generate comprehensive FAQ
    print("\nGenerating comprehensive FAQ...")
    faq_items = generate_comprehensive_faq(ranked, len(relevant))
    print(f"  - FAQ entries: {len(faq_items)}")
    
    # Save raw data for reference
    output_dir = repo_root / "docs"
    output_dir.mkdir(exist_ok=True)
    
    raw_path = output_dir / "faq_raw_data.json"
    with open(raw_path, 'w', encoding='utf-8') as f:
        json.dump({
            'ranked_qa': ranked[:50],
            'faq_items': faq_items,
            'stats': {
                'total_messages': len(messages),
                'relevant_messages': len(relevant),
                'qa_pairs_extracted': len(merged['qa_pairs']),
                'unique_qa': len(ranked)
            }
        }, f, ensure_ascii=False, indent=2)
    
    # Create DOCX
    docx_path = output_dir / "FAQ.docx"
    create_docx(faq_items, str(docx_path), len(relevant))
    
    # Also save markdown version
    md_path = output_dir / "faq_generated.md"
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write(f"# FAQ: Підтримка ArduPilot та StabX\n\n")
        
        categories = {}
        for item in faq_items:
            cat = item.get('category', 'Інше')
            if cat not in categories:
                categories[cat] = []
            categories[cat].append(item)
        
        for category, items in categories.items():
            f.write(f"## {category}\n\n")
            for item in items:
                f.write(f"**Q: {item.get('question', '')}**\n\n")
                f.write(f"A: {item.get('answer', '')}\n\n")
    
    print("\n" + "=" * 60)
    print("FAQ Generation Complete!")
    print("=" * 60)
    print(f"\nOutput files:")
    print(f"  - DOCX: {docx_path}")
    print(f"  - Markdown: {md_path}")
    print(f"  - Raw data: {raw_path}")


if __name__ == "__main__":
    main()
